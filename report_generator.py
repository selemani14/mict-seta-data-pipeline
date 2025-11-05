# report_generator.py (FINAL, VERIFIED CORRECTED VERSION)

import datetime
import os
import asyncio
from docx import Document
from typing import Any 

# Import from config to get logger and REPORT_DIRECTORY
# FIX: Removed the '.' to change to an ABSOLUTE import.
from config import logger, REPORT_DIRECTORY 

# --- II. SYNCHRONOUS HELPER ---

def _sync_generate_and_save_report(tool_name: str, params: dict[str, Any], file_path: str):
    """
    Performs all synchronous, blocking docx operations (creation, writing, saving)
    on a worker thread.
    """
    document = Document()
    document.add_heading(f'MICT SETA Tool Report: {tool_name.replace("_", " ").title()}', 0)
    document.add_paragraph(f'Report Generation Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    document.add_heading('Input Parameters', level=1)

    for key, value in params.items():
        if key not in ['logger']:
             document.add_paragraph(f'{key.replace("_", " ").title()}: {str(value)}')

    document.save(file_path)
    logger.info(f"[THREAD] Successfully saved report to: {file_path}")


# --- III. ASYNCHRONOUS WRAPPER ---

async def generate_word_report(tool_name: str, params: dict) -> dict:
    """
    The main report function that correctly offloads all synchronous I/O
    to a worker thread using asyncio.to_thread.
    """
    file_name = f"{tool_name}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(REPORT_DIRECTORY, file_name)

    logger.info(f"Preparing to generate Word report for tool: {tool_name}")

    try:
        await asyncio.to_thread(_sync_generate_and_save_report, tool_name, params, file_path)
    except Exception as e:
        logger.error(f"Failed to save document for {tool_name}: {e}")
        return {
            "status": "Report Generation Failed",
            "message": f"Server encountered an I/O error while saving report: {e}",
            "document_path": None,
        }

    return {
        "status": "Report Generated",
        "message": f"Successfully generated report for {tool_name}",
        "document_path": file_path,
    }
